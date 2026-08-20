"""
extract.py
----------
Extracts raw text from a CV/resume file (PDF or DOCX).

This is deliberately kept separate from parsing/scoring logic so each
piece can be tested and reused independently.
"""

from pathlib import Path
import pdfplumber
from docx import Document


class UnsupportedFileType(Exception):
    """Raised when the uploaded file isn't a PDF or DOCX."""
    pass


def extract_text_from_pdf(filepath: str) -> str:
    """Extract all text from a PDF file, page by page."""
    text_chunks = []
    with pdfplumber.open(filepath) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_chunks.append(page_text)
    return "\n".join(text_chunks)


def extract_text_from_docx(filepath: str) -> str:
    """Extract all text from a DOCX file, including tables."""
    doc = Document(filepath)
    text_chunks = []

    # Paragraphs (normal body text)
    for para in doc.paragraphs:
        if para.text.strip():
            text_chunks.append(para.text)

    # Tables (some resumes use tables for layout - ATS often chokes on these,
    # so we still extract the text but we'll flag table usage elsewhere)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_chunks.append(cell.text)

    return "\n".join(text_chunks)


def extract_text(filepath: str) -> str:
    """
    Dispatch to the correct extractor based on file extension.
    Returns raw extracted text as a single string.
    """
    ext = Path(filepath).suffix.lower()

    if ext == ".pdf":
        return extract_text_from_pdf(filepath)
    elif ext == ".docx":
        return extract_text_from_docx(filepath)
    else:
        raise UnsupportedFileType(
            f"Unsupported file type '{ext}'. Please upload a .pdf or .docx file."
        )


def has_docx_tables(filepath: str) -> bool:
    """
    Check whether a DOCX file uses tables for layout.
    ATS parsers frequently fail to correctly read text inside tables,
    so this is used later as a format-risk flag.
    """
    ext = Path(filepath).suffix.lower()
    if ext != ".docx":
        return False  # Table detection only implemented for DOCX for now
    doc = Document(filepath)
    return len(doc.tables) > 0


if __name__ == "__main__":
    import sys

    if len(sys.argv) != 2:
        print("Usage: python extract.py <path_to_cv>")
        sys.exit(1)

    text = extract_text(sys.argv[1])
    print(f"--- Extracted {len(text)} characters ---\n")
    print(text[:1000])  # preview first 1000 chars
